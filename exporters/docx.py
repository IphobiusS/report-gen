"""Exportacion a DOCX con python-docx puro (mismo diseno que el PDF)."""
import re
from pathlib import Path

import engine

from .common import sev_label, cvss_label, summary_sev


def _flatten_css_vars(css, extra):
    """LibreOffice no resuelve var(--x); sustituye cada variable por su valor."""
    varmap = dict(extra)
    for block in re.findall(r":root\s*{([^}]*)}", css):
        for m in re.finditer(r"--([\w-]+)\s*:\s*([^;]+);", block):
            varmap.setdefault("--" + m.group(1), m.group(2).strip())
    varmap.update(extra)

    def resolve(val, seen=()):  # resuelve variables que apuntan a otras
        def repl(m):
            name, fb = m.group(1), m.group(2)
            if name in varmap and name not in seen:
                return resolve(varmap[name], seen + (name,))
            return (fb.strip() if fb else "")
        return re.sub(r"var\((--[\w-]+)(?:,([^)]*))?\)", repl, val)

    resolved = {k: resolve(v) for k, v in varmap.items()}

    def repl(m):
        name, fb = m.group(1), m.group(2)
        return resolved.get(name, fb.strip() if fb else "")
    return re.sub(r"var\((--[\w-]+)(?:,([^)]*))?\)", repl, css)


def standalone_html(data, meta, L, engagement_dir):
    """HTML del informe con el CSS del tema embebido e imagenes como file://.
    Es el MISMO HTML que produce el PDF: asi el DOCX conserva el diseno."""
    env = engine.build_env()
    html = engine.render_html(env, data, engagement_dir, {}, False, L)
    theme = meta.get("theme", "serio")
    common = (engine.THEMES / "_common.css").read_text(encoding="utf-8")
    theme_css = (engine.THEMES / f"{theme}.css").read_text(encoding="utf-8")
    theme_css = theme_css.replace('@import "_common.css";', common)
    accent = (meta.get("branding") or {}).get("accent")
    theme_css = _flatten_css_vars(theme_css, {"--accent": accent} if accent else {})
    html = re.sub(r'<link rel="stylesheet"[^>]*>', f"<style>\n{theme_css}\n</style>", html, count=1)
    return html


def _palette(meta):
    theme = meta.get("theme", "serio")
    css = (engine.THEMES / f"{theme}.css").read_text(encoding="utf-8")
    common = (engine.THEMES / "_common.css").read_text(encoding="utf-8")
    css = css.replace('@import "_common.css";', common)
    varmap = {}
    for block in re.findall(r":root\s*{([^}]*)}", css):
        for m in re.finditer(r"--([\w-]+)\s*:\s*([^;]+);", block):
            varmap.setdefault("--" + m.group(1), m.group(2).strip())
    accent = (meta.get("branding") or {}).get("accent")
    if accent:
        varmap["--accent"] = accent

    def resolve(val, seen=()):
        def repl(m):
            n = m.group(1)
            if n in varmap and n not in seen:
                return resolve(varmap[n], seen + (n,))
            return (m.group(2).strip() if m.group(2) else "")
        return re.sub(r"var\((--[\w-]+)(?:,([^)]*))?\)", repl, val)

    res = {k: resolve(v).strip() for k, v in varmap.items()}

    def hx(name, default):
        v = res.get("--" + name, default).strip().lstrip("#")
        return v if len(v) == 6 and all(c in "0123456789abcdefABCDEF" for c in v) else default.lstrip("#")
    return hx


def _sev_chart_png(counts, hx, out):
    from PIL import Image, ImageDraw
    order = engine.SEVERITY_ORDER
    colmap = {s: hx("sev-" + s, "888888") for s in order}
    W, H = 900, 300
    im = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(im)
    maxc = max([counts[s] for s in order] + [1])
    bw, gap, x0, base, top = 90, 80, 60, 250, 40
    for i, s in enumerate(order):
        x = x0 + i * (bw + gap)
        h = int((base - top) * counts[s] / maxc)
        y = base - h
        rgb = tuple(int(colmap[s][j:j+2], 16) for j in (0, 2, 4))
        if counts[s] > 0:
            d.rectangle([x, y, x + bw, base], fill=rgb)
            d.text((x + bw/2 - 4, y - 22), str(counts[s]), fill=(20, 20, 30))
        else:
            d.line([x, base, x + bw, base], fill=(200, 200, 210), width=3)
        d.text((x + 6, base + 8), s.capitalize(), fill=(90, 90, 110))
    d.line([40, base, W - 20, base], fill=(210, 210, 220), width=2)
    im.save(out)


def to_docx(data, meta, L, engagement_dir, out_path):
    """DOCX con el diseno del tema (python-docx puro, sin dependencias externas).
    Portada, cabecera/pie con numero de pagina, cajas de finding, chips de
    severidad, bloques de codigo, grafico de severidades e imagenes."""
    import tempfile
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    lab = L["labels"]
    hx = _palette(meta)
    C = lambda name, d: RGBColor.from_string(hx(name, d))
    INK = C("ink", "0b0a1a"); INK_S = C("ink-soft", "2a2940"); MUTED = C("muted", "5f5e75")
    RULE = hx("rule", "d9d8e3"); PANEL = hx("panel", "f3f2f8")
    ACC = C("accent-ink", "1d8a12"); ACC_HEX = hx("accent", "39ff14")
    CODE_BG = hx("code-bg", "0b0a1a"); CODE_FG = C("code-inline-fg", "145f0c")
    CODE_INL_BG = hx("code-inline-bg", "eef7ee")
    SEV = {s: hx("sev-" + s, "888888") for s in engine.SEVERITY_ORDER}
    COVER_BG = hx("cover-bg", "0b0a1a"); COVER_FG = C("cover-fg", "ececf4")
    COVER_TITLE = C("cover-title", "ffffff"); COVER_SUB = C("cover-subtitle", "05d9e8")
    COVER_WM = C("cover-wordmark", "39ff14"); COVER_MUTED = C("cover-muted", "8f8fb0")
    COVER_RULE = hx("cover-rule", "1f5fa8"); COVER_TAG_FG = C("cover-tag-fg", "1a4e88")
    COVER_TAG_BORDER = hx("cover-tag-border", "1f5fa8"); COVER_META_BORDER = hx("cover-meta-border", "e2e5ec")

    MONO = "Consolas"
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def shd(el, fill):
        pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
        s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear"); s.set(qn("w:fill"), fill); pr.append(s)

    def cell_bg(cell, fill): shd(cell._tc, fill)
    def par_bg(p, fill): shd(p._p, fill)

    def par_border(p, color, edges=("bottom",), sz=6, space=2):
        pPr = p._p.get_or_add_pPr()
        bdr = OxmlElement("w:pBdr")
        for e in edges:
            b = OxmlElement("w:" + e)
            b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(sz))
            b.set(qn("w:space"), str(space)); b.set(qn("w:color"), color)
            bdr.append(b)
        pPr.append(bdr)

    def run_shade(run, fill):
        s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear"); s.set(qn("w:fill"), fill)
        run._r.get_or_add_rPr().append(s)

    def cell_border(cell, spec):
        tcPr = cell._tc.get_or_add_tcPr()
        tb = tcPr.find(qn("w:tcBorders"))
        if tb is None:
            tb = OxmlElement("w:tcBorders"); tcPr.append(tb)
        for edge, (color, sz) in spec.items():
            e = OxmlElement("w:" + edge)
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz)); e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
            tb.append(e)

    def run_border(run, color, sz=4, space=3):
        b = OxmlElement("w:bdr")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(sz)); b.set(qn("w:space"), str(space)); b.set(qn("w:color"), color)
        run._r.get_or_add_rPr().append(b)

    def field(p, instr):
        r = p.add_run()
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = instr
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
        r._r.append(b); r._r.append(t); r._r.append(e)

    def add_inline(p, text):
        for tok in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text or ""):
            if not tok:
                continue
            if tok.startswith("**") and tok.endswith("**"):
                p.add_run(tok[2:-2]).bold = True
            elif tok.startswith("`") and tok.endswith("`"):
                r = p.add_run(tok[1:-1]); r.font.name = MONO; r.font.size = Pt(9.5)
                r.font.color.rgb = CODE_FG; run_shade(r, CODE_INL_BG)
            else:
                p.add_run(tok)

    def add_markdown(text):
        for block in re.split(r"\n\s*\n", (text or "").strip()):
            block = block.strip()
            if not block:
                continue
            if all(ln.strip().startswith(("- ", "* ")) for ln in block.splitlines()):
                for ln in block.splitlines():
                    add_inline(doc.add_paragraph(style="List Bullet"), ln.strip()[2:])
            else:
                add_inline(doc.add_paragraph(), block.replace("\n", " "))

    def code_block(text):
        p = doc.add_paragraph(); par_bg(p, CODE_BG); par_border(p, ACC_HEX, ("left",), sz=18, space=4)
        p.paragraph_format.left_indent = Cm(0.2); p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        lines = text.strip().split("\n")
        for i, ln in enumerate(lines):
            if i:
                p.add_run().add_break()
            r = p.add_run(ln); r.font.name = MONO; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string("E9E9F2")
        return p

    sec_no = [0]

    def heading(text, level=1, sev=None):
        p = doc.add_heading("", level=level)
        for r in list(p.runs):
            r.text = ""
        if level == 1:
            sec_no[0] += 1
            n = p.add_run(str(sec_no[0]) + "  "); n.font.color.rgb = ACC; n.bold = True
        t = p.add_run(text); t.font.color.rgb = INK if level == 1 else INK_S
        t.bold = True
        if level == 1:
            t.font.size = Pt(15); par_border(p, hx("ink", "0b0a1a"), ("bottom",), sz=12)
        elif level == 2:
            t.font.size = Pt(12.5)
        else:
            t.font.size = Pt(10.5); t.font.color.rgb = INK_S
        if sev:
            sp = p.add_run("   " + sev_label(L, sev)); sp.bold = True
            sp.font.color.rgb = RGBColor.from_string("FFFFFF"); run_shade(sp, SEV.get(sev, "888888"))
        return p

    def kv_table(rows):
        t = doc.add_table(rows=0, cols=2); t.autofit = False
        t.columns[0].width = Cm(4.6); t.columns[1].width = Cm(12.4)
        for k, v in rows:
            cells = t.add_row().cells
            cells[0].width = Cm(4.6); cells[1].width = Cm(12.4)
            cell_bg(cells[0], PANEL)
            rk = cells[0].paragraphs[0].add_run(k); rk.font.size = Pt(8.5); rk.font.color.rgb = MUTED; rk.font.name = MONO
            add_inline(cells[1].paragraphs[0], v)
        doc.add_paragraph()
        return t

    def data_table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
        for i, hd in enumerate(headers):
            c = t.rows[0].cells[i]; cell_bg(c, PANEL)
            r = c.paragraphs[0].add_run(hd); r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = MUTED; r.font.name = MONO
        for row in rows:
            cs = t.add_row().cells
            for i, val in enumerate(row):
                add_inline(cs[i].paragraphs[0], str(val))
        doc.add_paragraph()

    def figure(step):
        fig = step.get("figure")
        if not fig or not fig.get("src"):
            return
        img = engagement_dir / fig["src"]
        if img.exists():
            try:
                doc.add_picture(str(img), width=Cm(15))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        cap = doc.add_paragraph()
        rc = cap.add_run(f"{L['figure']} {fig.get('number','')}. "); rc.bold = True; rc.italic = True; rc.font.size = Pt(9); rc.font.color.rgb = INK_S
        rr = cap.add_run(fig.get("caption", "")); rr.italic = True; rr.font.size = Pt(9); rr.font.color.rgb = MUTED

    def step(s):
        p = doc.add_paragraph()
        if s.get("lead"):
            rb = p.add_run(s["lead"] + " "); rb.bold = True
        if s.get("text_md"):
            add_inline(p, s["text_md"].replace("\n", " "))
        if s.get("command"):
            code_block(s["command"])
        figure(s)

    # ---------- Portada (seccion 0, sin cabecera/pie) ----------
    sec0 = doc.sections[0]
    sec0.page_width = Cm(21); sec0.page_height = Cm(29.7)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec0, m, Cm(0))
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = False
    tbl.columns[0].width = Cm(21)
    cell = tbl.rows[0].cells[0]; cell.width = Cm(21); cell_bg(cell, COVER_BG)
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight"); trH.set(qn("w:val"), str(int(29.7 * 566.9))); trH.set(qn("w:hRule"), "exact"); trPr.append(trH)
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "start", "end", "bottom"):
        el = OxmlElement("w:" + side); el.set(qn("w:w"), str(int(2.2 * 566.9))); el.set(qn("w:type"), "dxa"); tcMar.append(el)
    cell._tc.get_or_add_tcPr().append(tcMar)
    br = meta.get("branding") or {}
    CW = 16.6  # ancho de contenido de la portada (cm)

    def cover_gap(n=1):
        for _ in range(n):
            gp = cell.add_paragraph(); gp.paragraph_format.space_after = Pt(0)
            gp.add_run(" ").font.size = Pt(11)

    def accent_rule():
        t = cell.add_table(rows=1, cols=1); t.autofit = False
        t.columns[0].width = Cm(4.6); c = t.rows[0].cells[0]; c.width = Cm(4.6)
        cell_border(c, {"bottom": (COVER_RULE, 18)})
        cp = c.paragraphs[0]; cp.paragraph_format.space_after = Pt(0); cp.paragraph_format.space_before = Pt(0)
        cp.add_run(" ").font.size = Pt(2)

    # Cabecera: wordmark izquierda + tag CONFIDENCIAL en caja a la derecha, con regla inferior
    hp = cell.paragraphs[0]
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(CW), WD_TAB_ALIGNMENT.RIGHT)
    rw = hp.add_run(br.get("wordmark", "iphobiuss")); rw.bold = True; rw.font.size = Pt(16); rw.font.color.rgb = COVER_WM
    hp.add_run("\t")
    rt = hp.add_run(" " + br.get("confidential_text", "CONFIDENTIAL") + " ")
    rt.font.size = Pt(8.5); rt.font.color.rgb = COVER_TAG_FG; run_border(rt, COVER_TAG_BORDER, sz=4, space=4)
    par_border(hp, COVER_RULE, ("bottom",), sz=8, space=8)

    cover_gap(11)
    accent_rule()
    tp = cell.add_paragraph(); tp.paragraph_format.space_before = Pt(6)
    rt2 = tp.add_run(meta.get("report_title", "")); rt2.bold = True; rt2.font.size = Pt(26); rt2.font.color.rgb = COVER_TITLE
    if meta.get("report_subtitle"):
        sp = cell.add_paragraph(); rs = sp.add_run(meta["report_subtitle"]); rs.font.size = Pt(12); rs.font.color.rgb = COVER_SUB
    accent_rule()

    cover_gap(10)
    mt = cell.add_table(rows=0, cols=2); mt.autofit = False
    mt.columns[0].width = Cm(4.2); mt.columns[1].width = Cm(12.4)
    for k, v in [(lab["client"], meta.get("client", "")),
                 (lab["cover_assessor"], f"{meta.get('assessor','')}" + (f" · {meta['assessor_title']}" if meta.get("assessor_title") else "")),
                 (lab["cover_date"], meta.get("date", "")), (lab["cover_version"], meta.get("version", ""))]:
        if not v:
            continue
        cells = mt.add_row().cells
        cells[0].width = Cm(4.2); cells[1].width = Cm(12.4)
        cell_border(cells[0], {"bottom": (COVER_META_BORDER, 4)})
        cell_border(cells[1], {"bottom": (COVER_META_BORDER, 4)})
        rk = cells[0].paragraphs[0].add_run(k); rk.font.size = Pt(9); rk.font.color.rgb = COVER_MUTED; rk.font.name = MONO
        rv = cells[1].paragraphs[0].add_run(v); rv.font.size = Pt(10); rv.font.color.rgb = COVER_FG


    # ---------- Cuerpo (seccion 1, con cabecera/pie) ----------
    doc.add_section(WD_SECTION.NEW_PAGE)
    sec1 = doc.sections[-1]
    sec1.page_width = Cm(21); sec1.page_height = Cm(29.7)
    sec1.left_margin = sec1.right_margin = Cm(1.8); sec1.top_margin = Cm(1.6); sec1.bottom_margin = Cm(1.6)
    conf = br.get("confidential_text", "CONFIDENTIAL")
    rhright = f"{conf}  |  v{meta.get('version','')}  |  {meta.get('date','')}"
    sec1.header.is_linked_to_previous = False
    hp = sec1.header.paragraphs[0]; hp.paragraph_format.tab_stops.add_tab_stop(Cm(17.4), WD_TAB_ALIGNMENT.RIGHT)
    r = hp.add_run(br.get("wordmark", "iphobiuss")); r.bold = True; r.font.color.rgb = ACC; r.font.size = Pt(8); r.font.name = MONO
    r = hp.add_run("\t" + rhright); r.font.size = Pt(7.5); r.font.color.rgb = MUTED; r.font.name = MONO
    par_border(hp, RULE, ("bottom",), sz=4)
    sec1.footer.is_linked_to_previous = False
    fp = sec1.footer.paragraphs[0]; fp.paragraph_format.tab_stops.add_tab_stop(Cm(17.4), WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run(br.get("byline") or meta.get("assessor", "")); r.font.size = Pt(7.5); r.font.color.rgb = MUTED; r.font.name = MONO
    r = fp.add_run("\t" + L["running"]["page"] + " "); r.font.size = Pt(7.5); r.font.color.rgb = MUTED; r.font.name = MONO
    field(fp, "PAGE")
    r = fp.add_run(" " + L["running"]["of"] + " "); r.font.size = Pt(7.5); r.font.color.rgb = MUTED; r.font.name = MONO
    field(fp, "NUMPAGES")

    def newpage():
        doc.add_page_break()

    def render_summary_body():
        sumf = engine.summary_findings(data["findings"]); counts = engine.severity_counts(data["findings"])
        parts = [f"{counts[s]} {sev_label(L, s).lower()}" for s in engine.SEVERITY_ORDER if counts[s] > 0]
        total = len(sumf); noun = L["summary"]["finding"] if total == 1 else L["summary"]["findings"]
        doc.add_paragraph(f"{L['summary']['identified']} {total} {noun} {L['summary']['rating']} {', '.join(parts) if parts else L['summary']['none']}.")
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
            chart = tf.name
        _sev_chart_png(counts, hx, chart)
        doc.add_picture(chart, width=Cm(9)); Path(chart).unlink(missing_ok=True)
        if sumf:
            rows = [[i, summary_sev(L, f), f.get("title", "")] for i, f in enumerate(sumf, 1)]
            data_table([lab["col_num"], lab["col_sev"], lab["col_finding"]], rows)

    def render_one_finding(f):
        if f.get("mode") == "machine":
            host = f.get("host") or {}
            heading(f"{f.get('id','')} {host.get('name') or f.get('title','')}", 2)
            rows = []
            if host.get("ip"): rows.append((lab["ip"], host["ip"]))
            if f.get("open_ports"): rows.append((lab["open_ports"], f["open_ports"]))
            if host.get("os"): rows.append((lab["os"], host["os"]))
            if rows: kv_table(rows)
            if f.get("summary_md"):
                p = doc.add_paragraph(); rb = p.add_run(lab["attack_path"] + ". "); rb.bold = True
                add_inline(p, f["summary_md"].replace("\n", " "))
            for ph in f.get("phases", []):
                heading(ph.get("name", ""), 3)
                for s in ph.get("steps", []): step(s)
            if f.get("proof"):
                heading(lab["proof"], 3)
                data_table([lab["file"], lab["hash"]], [[p.get("name", ""), p.get("value", "")] for p in f["proof"]])
        else:
            heading(f"{f.get('id','')} {f.get('title','')}", 2, sev=f.get("severity"))
            rows = []
            if f.get("cwe"): rows.append((lab["cwe"], f["cwe"]))
            if f.get("cvss_vector") or f.get("cvss"): rows.append((cvss_label(lab, f), (f.get("cvss", "") + "  " + f.get("cvss_vector", "")).strip()))
            if f.get("affected"): rows.append((lab["affected"], f["affected"]))
            if rows: kv_table(rows)
            for key, lbl in [("description_md", lab["desc_root"]), ("impact_md", lab["impact"]), ("remediation_md", lab["remediation"])]:
                if f.get(key):
                    p = doc.add_paragraph(); rb = p.add_run(lbl + ". "); rb.bold = True
                    add_inline(p, f[key].replace("\n", " "))
            if f.get("references"):
                heading(lab["references"], 3)
                for r in f["references"]: doc.add_paragraph(r, style="List Bullet")
            if f.get("walkthrough"):
                heading(lab["walkthrough"], 3)
                for s in f["walkthrough"]: step(s)
            rs = f.get("remediation_summary") or {}
            if any(rs.values()):
                heading(lab["rem_summary"], 3)
                for key, lbl in [("short_md", lab["short_term"]), ("medium_md", lab["medium_term"]), ("long_md", lab["long_term"])]:
                    if rs.get(key):
                        p = doc.add_paragraph(style="List Bullet"); rb = p.add_run(lbl + ": "); rb.bold = True
                        add_inline(p, rs[key])

    # ---- Rama por secciones (mismo orden/contenido que el PDF) ----
    if (data.get("report") or {}).get("sections"):
        import sections as _sections
        rsecs = _sections.resolve_sections(data, L, meta.get("lang", "en"))

        def render_generic(s):
            vals = [f for f in s["fields"] if f.get("value")]
            multi = len(vals) > 1
            for f in s["fields"]:
                v = f.get("value")
                if v:
                    if multi: heading(f["label"], 3)
                    if f["type"] == "markdown": add_markdown(v)
                    elif f["type"] == "text": doc.add_paragraph(v)
                    elif f["type"] == "table":
                        data_table([c["label"] for c in f.get("columns", [])], [[row.get(c["key"], "") for c in f.get("columns", [])] for row in v])
                    elif f["type"] == "list":
                        for it in v: doc.add_paragraph(str(it), style="List Bullet")
                elif f.get("default_lang"):
                    client = (meta.get("client") or lab["client"]).rstrip(".")
                    for para in L["prose"][f["default_lang"]]:
                        add_markdown(para.replace("{client}", client).replace("{assessor}", meta.get("assessor", "")))

        for i, s in enumerate(rsecs, 1):
            if i > 1: newpage()
            heading(s["title"], 1)
            sp = s.get("special")
            if sp == "contacts":
                c = meta.get("contacts") or {}
                for grp, gt in [("client", lab["client"]), ("assessor", lab["team"])]:
                    ppl = c.get(grp) or []
                    if ppl:
                        heading(gt, 2)
                        data_table([lab["name"], lab["title"], lab["email"]], [[p.get("name", ""), p.get("title", ""), p.get("email", "")] for p in ppl])
            elif sp == "scope":
                render_generic(s)
                if meta.get("scope"):
                    heading(lab["scope"], 2)
                    data_table([lab["target"], lab["description"]], [[x.get("target", ""), x.get("description", "")] for x in meta["scope"]])
            elif sp == "summary":
                render_summary_body()
            elif sp == "findings":
                for idx, f in enumerate(data["findings"]):
                    if idx: newpage()
                    render_one_finding(f)
            elif sp == "appendix":
                render_generic(s)
                appx = engine.appendix_rows(data["findings"])
                if appx:
                    heading(lab["appendix"], 2)
                    doc.add_paragraph(L.get("appendix_intro", ""))
                    data_table([lab["host"], lab["item"], lab["value"], lab["notes"]], [[r["host"], r["item"], r["value"], r["notes"]] for r in appx])
            else:
                render_generic(s)
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return out_path

    # Confidencialidad
    heading(lab["confidentiality"], 1)
    if meta.get("confidentiality_md"):
        add_markdown(meta["confidentiality_md"])
    else:
        client = (meta.get("client") or lab["client"]).rstrip(".")
        for para in L["prose"]["confidentiality"]:
            add_markdown(para.replace("{client}", client))

    # Contactos
    contacts = meta.get("contacts") or {}
    if contacts.get("client") or contacts.get("assessor"):
        newpage(); heading(lab["contacts"], 1)
        for grp, gt in [("client", lab["client"]), ("assessor", lab["team"])]:
            ppl = contacts.get(grp) or []
            if ppl:
                heading(gt, 2)
                data_table([lab["name"], lab["title"], lab["email"]], [[p.get("name", ""), p.get("title", ""), p.get("email", "")] for p in ppl])

    # Overview
    newpage(); heading(lab["overview"], 1); heading(lab["approach"], 2)
    if meta.get("overview", {}).get("approach_md"):
        add_markdown(meta["overview"]["approach_md"])
    else:
        client = (meta.get("client") or lab["client"]).rstrip(".")
        for para in L["prose"]["approach"]:
            add_markdown(para.replace("{assessor}", meta.get("assessor", "")).replace("{client}", client))
    if meta.get("scope"):
        heading(lab["scope"], 2)
        data_table([lab["target"], lab["description"]], [[s.get("target", ""), s.get("description", "")] for s in meta["scope"]])

    # Summary
    sumf = engine.summary_findings(data["findings"]); counts = engine.severity_counts(data["findings"])
    newpage(); heading(lab["summary"], 1)
    parts = [f"{counts[s]} {sev_label(L, s).lower()}" for s in engine.SEVERITY_ORDER if counts[s] > 0]
    total = len(sumf); noun = L["summary"]["finding"] if total == 1 else L["summary"]["findings"]
    doc.add_paragraph(f"{L['summary']['identified']} {total} {noun} {L['summary']['rating']} {', '.join(parts) if parts else L['summary']['none']}.")
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
        chart = tf.name
    _sev_chart_png(counts, hx, chart)
    doc.add_picture(chart, width=Cm(9)); Path(chart).unlink(missing_ok=True)
    if sumf:
        rows = [[i, summary_sev(L, f), f.get("title", "")] for i, f in enumerate(sumf, 1)]
        data_table([lab["col_num"], lab["col_sev"], lab["col_finding"]], rows)

    # Findings
    newpage(); heading(lab["findings"], 1)
    for idx, f in enumerate(data["findings"]):
        if idx:
            newpage()
        if f.get("mode") == "machine":
            host = f.get("host") or {}
            heading(f"{f.get('id','')} {host.get('name') or f.get('title','')}", 2)
            rows = []
            if host.get("ip"):
                rows.append((lab["ip"], host["ip"]))
            if f.get("open_ports"):
                rows.append((lab["open_ports"], f["open_ports"]))
            if host.get("os"):
                rows.append((lab["os"], host["os"]))
            if rows:
                kv_table(rows)
            if f.get("summary_md"):
                p = doc.add_paragraph(); rb = p.add_run(lab["attack_path"] + ". "); rb.bold = True
                add_inline(p, f["summary_md"].replace("\n", " "))
            for ph in f.get("phases", []):
                heading(ph.get("name", ""), 3)
                for s in ph.get("steps", []):
                    step(s)
            if f.get("proof"):
                heading(lab["proof"], 3)
                data_table([lab["file"], lab["hash"]], [[p.get("name", ""), p.get("value", "")] for p in f["proof"]])
        else:
            heading(f"{f.get('id','')} {f.get('title','')}", 2, sev=f.get("severity"))
            rows = []
            if f.get("cwe"):
                rows.append((lab["cwe"], f["cwe"]))
            if f.get("cvss_vector") or f.get("cvss"):
                rows.append((cvss_label(lab, f), (f.get("cvss", "") + "  " + f.get("cvss_vector", "")).strip()))
            if f.get("affected"):
                rows.append((lab["affected"], f["affected"]))
            if rows:
                kv_table(rows)
            for key, lbl in [("description_md", lab["desc_root"]), ("impact_md", lab["impact"]), ("remediation_md", lab["remediation"])]:
                if f.get(key):
                    p = doc.add_paragraph(); rb = p.add_run(lbl + ". "); rb.bold = True
                    add_inline(p, f[key].replace("\n", " "))
            if f.get("references"):
                heading(lab["references"], 3)
                for r in f["references"]:
                    doc.add_paragraph(r, style="List Bullet")
            if f.get("walkthrough"):
                heading(lab["walkthrough"], 3)
                for s in f["walkthrough"]:
                    step(s)
            rs = f.get("remediation_summary") or {}
            if any(rs.values()):
                heading(lab["rem_summary"], 3)
                for key, lbl in [("short_md", lab["short_term"]), ("medium_md", lab["medium_term"]), ("long_md", lab["long_term"])]:
                    if rs.get(key):
                        p = doc.add_paragraph(style="List Bullet"); rb = p.add_run(lbl + ": "); rb.bold = True
                        add_inline(p, rs[key])

    # Appendix
    appx = engine.appendix_rows(data["findings"])
    if appx:
        newpage(); heading(lab["appendix"], 1)
        doc.add_paragraph(L.get("appendix_intro", ""))
        data_table([lab["host"], lab["item"], lab["value"], lab["notes"]], [[r["host"], r["item"], r["value"], r["notes"]] for r in appx])

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
