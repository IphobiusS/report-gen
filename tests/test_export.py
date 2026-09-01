"""Tests de exportacion: Markdown y DOCX conscientes de las secciones activas."""
import tempfile
from pathlib import Path

import engine
import export

SECTION_DATA = {
    "meta": {"lang": "es", "report_title": "T", "theme": "corporativo", "client": "Acme"},
    "report": {"sections": [
        {"key": "confidentiality"},
        {"key": "executive_summary", "summary": "Resumen ejecutivo de prueba."},
        {"key": "findings_summary"},
        {"key": "findings"},
        {"key": "remediation_plan", "short_term": "Parche."},
        {"key": "appendix", "tools_used": "nmap, netexec"},
    ]},
    "findings": [
        {"id": "F1", "mode": "vuln", "title": "RCE", "severity": "critical",
         "cvss": "9.8", "description_md": "Texto `x`."},
    ],
}


def test_markdown_section_order():
    L = engine.load_lang("es")
    md = export.to_markdown(SECTION_DATA, SECTION_DATA["meta"], L, Path("."))
    headings = [ln for ln in md.splitlines() if ln.startswith("## ")]
    titles = [h.split(". ", 1)[1] for h in headings]
    assert titles == [
        "Declaración de confidencialidad", "Resumen ejecutivo",
        "Resumen de hallazgos", "Hallazgos", "Plan de remediación", "Apéndices",
    ]
    assert "nmap, netexec" in md
    assert "F1" in md


def test_docx_section_aware_opens_and_orders():
    from docx import Document
    L = engine.load_lang("es")
    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "r.docx"
        export.to_docx(SECTION_DATA, SECTION_DATA["meta"], L, Path(d).resolve(), out)
        assert out.exists() and out.stat().st_size > 5000
        doc = Document(str(out))
        h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        # deben aparecer las secciones activas como H1
        joined = " | ".join(h1)
        for t in ("Resumen ejecutivo", "Hallazgos", "Plan de remediación", "Apéndices"):
            assert t in joined
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "nmap, netexec" in full


def test_markdown_fixed_model_still_works():
    L = engine.load_lang("es")
    data = {"meta": {"lang": "es", "report_title": "T", "theme": "serio"},
            "findings": [{"id": "F1", "mode": "vuln", "title": "X", "severity": "high"}]}
    md = export.to_markdown(data, data["meta"], L, Path("."))
    assert "F1" in md


def test_export_with_machine_finding():
    """Exportar un informe con un hallazgo 'machine' (sin severity) no debe crashear
    en DOCX ni MD; la maquina aparece en el resumen como 'Maquina' sin score."""
    import tempfile, os
    from pathlib import Path
    import export, engine
    data = {"meta": {"lang": "es", "report_title": "M", "theme": "serio", "branding": {}},
            "report": {"sections": [{"key": "findings_summary"}, {"key": "findings"}]},
            "findings": [
                {"id": "F1", "mode": "vuln", "title": "V", "severity": "high", "cvss": "7.5"},
                {"id": "M1", "mode": "machine", "title": "Box", "host": {"name": "Box", "ip": "10.0.0.1"},
                 "phases": [{"name": "Init", "steps": [{"lead": "x", "command": "whoami"}]}],
                 "proof": [{"name": "user.txt", "value": "abc"}]}]}
    L = engine.load_lang("es")
    md = export.to_markdown(data, data["meta"], L, Path("."))
    assert "Box" in md and "Init" in md and "whoami" in md
    d = tempfile.mkdtemp()
    export.to_docx(data, data["meta"], L, Path(d).resolve(), os.path.join(d, "m.docx"))
    from docx import Document
    doc = Document(os.path.join(d, "m.docx"))
    txt = "\n".join(p.text for p in doc.paragraphs) + "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Box" in txt
